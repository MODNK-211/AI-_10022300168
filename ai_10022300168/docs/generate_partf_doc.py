from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parent
DOCX_PATH = ROOT / "PART_F_Architecture_System_Design.docx"
IMG_PATH = ROOT / "part_f_data_flow.png"


def draw_data_flow_image(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    boxes = {
        "sources": (0.08, 0.82, "Data Sources\nCSV + PDF"),
        "ingest": (0.30, 0.82, "Data Loader\n(clean + normalize)"),
        "chunk": (0.52, 0.82, "Chunker\n(fixed/sentence)"),
        "embed": (0.74, 0.82, "Embedder\n(all-MiniLM-L6-v2)"),
        "faiss": (0.90, 0.82, "Vector Store\nFAISS + metadata"),
        "query": (0.08, 0.48, "User Query\n(Streamlit UI)"),
        "retrieve": (0.32, 0.48, "Hybrid Retriever\nFAISS + TF-IDF + alpha"),
        "prompt": (0.56, 0.48, "Prompt Builder\n(top snippets + citations)"),
        "llm": (0.78, 0.48, "LLM Client\nGroq (HF fallback)"),
        "resp": (0.92, 0.48, "Grounded Response\nshown to user"),
        "fb": (0.50, 0.18, "Feedback Loop\n(thumbs up/down boosts)"),
    }

    for _, (x, y, label) in boxes.items():
        ax.text(
            x,
            y,
            label,
            ha="center",
            va="center",
            fontsize=9,
            bbox=dict(boxstyle="round,pad=0.35", fc="#f4f8ff", ec="#2f5597", lw=1.2),
        )

    arrows = [
        ("sources", "ingest"),
        ("ingest", "chunk"),
        ("chunk", "embed"),
        ("embed", "faiss"),
        ("query", "retrieve"),
        ("faiss", "retrieve"),
        ("retrieve", "prompt"),
        ("prompt", "llm"),
        ("llm", "resp"),
        ("resp", "query"),
        ("retrieve", "fb"),
        ("fb", "retrieve"),
    ]

    for a, b in arrows:
        x1, y1, _ = boxes[a]
        x2, y2, _ = boxes[b]
        ax.annotate(
            "",
            xy=(x2, y2),
            xytext=(x1, y1),
            arrowprops=dict(arrowstyle="->", lw=1.4, color="#1f4e79"),
        )

    ax.set_title("The Acity Oracle: Data Flow and Component Interaction", fontsize=13, pad=16)
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def build_docx(docx_path: Path, image_path: Path) -> None:
    doc = Document()
    title = doc.add_heading("PART F: Architecture & System Design", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "System: The Acity Oracle (RAG chatbot for Ghana Election Results and Ghana 2025 Budget data)."
    )

    doc.add_heading("1. Architecture Overview", level=1)
    doc.add_paragraph(
        "The system has three layers: (a) data ingestion/indexing, (b) query-time RAG pipeline, "
        "and (c) Streamlit UI with a feedback loop. Data is prepared once and reused; each user query "
        "runs through retrieval, prompt construction, and model inference."
    )

    doc.add_heading("2. Data Flow Drawing", level=1)
    doc.add_paragraph("Figure 1 shows the end-to-end data flow and feedback loop.")
    doc.add_picture(str(image_path), width=Inches(6.8))
    cap = doc.add_paragraph("Figure 1: End-to-end data flow and components interaction.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("3. Components Interaction", level=1)
    interactions = [
        "data_loader.py cleans CSV/PDF data before chunk creation.",
        "chunker.py splits documents using fixed-size or sentence-aware strategy.",
        "embedder.py converts chunks into vectors; vector_store.py persists FAISS index and metadata.",
        "retriever.py combines semantic FAISS results and TF-IDF keyword matching using alpha-weighted fusion.",
        "prompt_builder.py selects highest-ranked snippets and builds a citation-oriented prompt.",
        "llm_client.py sends prompt to Groq (or Hugging Face fallback) and returns grounded response.",
        "app.py displays response, retrieved context, logs, and captures thumbs up/down feedback.",
        "feedback.py stores chunk-level boosts and injects them into future retrieval ranking.",
    ]
    for item in interactions:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. Why This Design Fits the Domain", level=1)
    suitability = [
        "Factual grounding is essential for election and budget questions; RAG reduces hallucinations by forcing source-based context.",
        "The domain has mixed source types (structured CSV + narrative PDF), so dual chunking strategies improve retrieval quality.",
        "Hybrid retrieval captures both semantic meaning and exact domain terms such as party acronyms and budget entities.",
        "Transparent UI (chunks, scores, prompt, logs) supports academic traceability and easier marking.",
        "The architecture is lightweight and practical for student deployment: FAISS on CPU plus API-hosted LLM.",
    ]
    for item in suitability:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(docx_path)


if __name__ == "__main__":
    draw_data_flow_image(IMG_PATH)
    build_docx(DOCX_PATH, IMG_PATH)
    print(f"Created: {DOCX_PATH}")
    print(f"Created: {IMG_PATH}")
